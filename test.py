from docx import Document
import os
from docx.shared import Cm
from PIL import Image

def get_pic_list(cur_dir):
    pic_list = []
    for path, dirs, files in os.walk(cur_dir):
        for file in files:
            if os.path.splitext(file)[1].lower() == '.jpg':
                pic_list.append(file)
            elif os.path.splitext(file)[1].lower() == '.png':
                pic_list.append(file)
            elif os.path.splitext(file)[1].lower() == '.gif':
                pic_list.append(file)
    return pic_list



if __name__ == "__main__":
    path = 'C:\KakaoTalk_Chats_2017-02-12_23.00.56'
    filename = path + '/KakaoTalkChats.txt'
    pic_list = get_pic_list(path)

    document = Document()

    with open(filename, 'r',-1, 'utf-8') as f:
        lines = f.readlines()
        for line in lines:
            find=False
            for p_list in pic_list:
                if p_list in line:
                    #pic_list.remove(p_list)
                    line = line.replace(p_list,"")
                    line = line.replace('\n',"")
                    paragraph = document.add_paragraph(line)
                    with Image.open(path + '/' + p_list) as im:
                        p_width, p_height = im.size
                    if p_width <= 15.93 and p_height <= 15.93:
                        paragraph = document.add_picture(path + '/' + p_list, width=Cm(p_width))
                    elif p_width > 15.93 and p_height < 15.93:  # width > height
                        paragraph = document.add_picture(path + '/' + p_list, width=Cm(15.93))
                    elif p_width < 15.93 and p_height > 15.93: # width < height
                        paragraph = document.add_picture(path + '/' + p_list, height=Cm(15.93))
                    elif p_width > 15.93 and p_height > 15.93:
                        if p_width > p_height:
                            paragraph = document.add_picture(path + '/' + p_list, width=Cm(15.93))
                        elif p_width < p_height:
                            paragraph = document.add_picture(path + '/' + p_list, height=Cm(15.93))

                    find=True
                    break
            if find == True:
                find = False
            else:
                paragraph = document.add_paragraph(line)
        document.save(path + '/demo.docx')